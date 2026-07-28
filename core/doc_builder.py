import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import subprocess
import copy


DEFAULT_PROFILE = "Profesional proactivo y capacitado con amplia experiencia orientada a resultados."
DEFAULT_EXPERIENCE_BLOCKS = []
DEFAULT_SKILLS_LINES = []
DEFAULT_EDUCATION_LINES = []

# ─────────────────────────────────────────────────────────────────────────────
# UTILIDADES BÁSICAS
# ─────────────────────────────────────────────────────────────────────────────

def _all_paragraphs(doc):
    """Devuelve todos los párrafos del doc (body + celdas de tabla)."""
    paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras.extend(cell.paragraphs)
    return paras


def clean_text(text):
    cleaned = text.strip().lstrip("●•-* \t0123456789.")
    return cleaned.strip()


def replace_placeholder(doc, placeholder, replacement_text):
    """
    Reemplaza {{PLACEHOLDER}} en el doc preservando el formato del run.
    """
    for para in _all_paragraphs(doc):
        if placeholder not in para.text:
            continue
        full_text = "".join(r.text for r in para.runs)
        if placeholder not in full_text:
            continue
        new_text = full_text.replace(placeholder, replacement_text)
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""


# ─────────────────────────────────────────────────────────────────────────────
# LECTURA DE VIÑETAS DEL DOCX (para referencia de la IA)
# ─────────────────────────────────────────────────────────────────────────────

def read_docx_bullets(template_path):
    """
    Lee todas las líneas de texto del docx (excluye placeholders y headings cortos).
    Devuelve lista de dicts con {text, clean}.
    """
    if not os.path.exists(template_path):
        return []
    bullets = []
    try:
        doc = docx.Document(template_path)
        for para in _all_paragraphs(doc):
            text = para.text.strip()
            if not text or (text.startswith("{{") and text.endswith("}}")):
                continue
            clean = clean_text(text)
            if len(clean) < 5:
                continue
            bullets.append({"text": text, "clean": clean})
    except Exception as e:
        print(f"[!] Error leyendo viñetas del docx: {e}")
    return bullets


# ─────────────────────────────────────────────────────────────────────────────
# INYECCIÓN DINÁMICA DE EXPERIENCIA  ({{EXPERIENCE_SECTION}})
# ─────────────────────────────────────────────────────────────────────────────

def _find_placeholder_para(doc, placeholder):
    """Devuelve el párrafo que contiene el placeholder, o None."""
    for para in _all_paragraphs(doc):
        if placeholder in para.text:
            return para
    return None


def _copy_rpr(src_run):
    """Clona el elemento w:rPr de un run dado."""
    rpr = src_run._element.find(qn('w:rPr'))
    return copy.deepcopy(rpr) if rpr is not None else None


def _find_ref_bullet_para(doc):
    """Devuelve el primer párrafo de viñeta real del doc para copiar su pPr y rPr."""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and (text.startswith("●") or text.startswith("•")):
            return para
    return None


def _get_ref_bullet_info(doc):
    """
    Lee el primer párrafo de viñeta del template y devuelve:
    - ref_rpr: elemento w:rPr (fuente, tamaño)
    - ref_ppr: elemento w:pPr COMPLETO (estilo + sangría + spacing del template)
    """
    ref_para = _find_ref_bullet_para(doc)
    ref_rpr = None
    ref_ppr = None
    if ref_para:
        if ref_para.runs:
            for run in ref_para.runs:
                if run.text.strip():
                    ref_rpr = run._element.find(qn('w:rPr'))
                    break
        pPr = ref_para._element.find(qn('w:pPr'))
        if pPr is not None:
            ref_ppr = copy.deepcopy(pPr)   # copia completa: estilo, ind, numPr, etc.
    return ref_rpr, ref_ppr


def _set_paragraph_spacing(pPr, before_twips=0, after_twips=0, line_twips=276):
    """
    Fuerza el interlineado y espaciado del párrafo.
    line_twips=276 → 1.15× (240=simple, 276=1.15, 360=1.5)
    """
    # Eliminar cualquier w:spacing previo
    for existing in pPr.findall(qn('w:spacing')):
        pPr.remove(existing)
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before_twips))
    spacing.set(qn('w:after'), str(after_twips))
    spacing.set(qn('w:line'), str(line_twips))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def _make_paragraph(text, bold=False, italic=False, ref_rpr=None,
                    is_bullet=False, before_twips=0, after_twips=0,
                    font_size_half_pts=None, color_hex=None,
                    ref_ppr=None,
                    indent_left=0, indent_hanging=0):
    """
    Crea un w:p con el texto y formato dado.
    - ref_ppr: si se provee, se usa como base del pPr (hereda estilo y sangría del template).
              El w:spacing se sobreescribe siempre para garantizar interlineado 1.15.
    - font_size_half_pts: sobreescribe tamaño de fuente (22=11pt, 20=10pt)
    - color_hex: color en hex sin '#', ej '595959' para gris
    - indent_left / indent_hanging: solo se usan si ref_ppr es None
    """
    p = OxmlElement('w:p')

    # --- pPr ---
    if ref_ppr is not None:
        # Copiar el pPr completo del template (estilo + sangría + numPr)
        pPr = copy.deepcopy(ref_ppr)
        # Sobreescribir SOLO el spacing para garantizar interlineado ajustado
        _set_paragraph_spacing(pPr, before_twips=before_twips, after_twips=after_twips)
    else:
        pPr = OxmlElement('w:pPr')
        _set_paragraph_spacing(pPr, before_twips=before_twips, after_twips=after_twips)
        if indent_left or indent_hanging:
            ind = OxmlElement('w:ind')
            if indent_left:
                ind.set(qn('w:left'), str(indent_left))
            if indent_hanging:
                ind.set(qn('w:hanging'), str(indent_hanging))
            pPr.append(ind)
    
    # Configuración de sangría francesa (Hanging Indent) para viñetas
    if is_bullet:
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            ind = OxmlElement('w:ind')
            pPr.append(ind)
        # left=360 twips (0.25 in), hanging=360 twips (0.25 in)
        # La viñeta (●) queda a 0" y todo el texto (línea 1 y líneas siguientes) se alinea a 0.25"
        ind.set(qn('w:left'), '800')
        ind.set(qn('w:hanging'), '200')

    p.append(pPr)

    # --- rPr ---
    r = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')

    if ref_rpr is not None:
        for tag in [qn('w:rFonts')]:
            el = ref_rpr.find(tag)
            if el is not None:
                rpr.append(copy.deepcopy(el))

    if font_size_half_pts:
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(font_size_half_pts))
        rpr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(font_size_half_pts))
        rpr.append(szCs)
    elif ref_rpr is not None:
        for tag in [qn('w:sz'), qn('w:szCs')]:
            el = ref_rpr.find(tag)
            if el is not None:
                rpr.append(copy.deepcopy(el))

    if color_hex:
        color_el = OxmlElement('w:color')
        color_el.set(qn('w:val'), color_hex)
        rpr.append(color_el)

    if bold:
        rpr.append(OxmlElement('w:b'))
    if italic:
        rpr.append(OxmlElement('w:i'))
        rpr.append(OxmlElement('w:iCs'))

    r.append(rpr)

    # Para viñeta con sangría francesa en Word: el prefijo DEBE llevar un tabulador '\t'
    # Así la 1ra línea salta al 0.25" y la 2da línea envuelve alineada al 0.25"
    prefix = "●\t" if is_bullet else ""

    t = OxmlElement('w:t')
    t.text = prefix + text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p


def inject_experience_section(doc, experience_blocks):
    """
    Busca {{EXPERIENCE_SECTION}} y reemplaza con bloques de experiencia formateados.
    Los bullets heredan el estilo COMPLETO del template (pPr) para que la sangría
    sea idéntica a los bullets de LANGUAGES.
    """
    placeholder_para = _find_placeholder_para(doc, "{{EXPERIENCE_SECTION}}")
    if placeholder_para is None:
        print("[!] No se encontró {{EXPERIENCE_SECTION}} en el docx.")
        print("    → Abrí el .docx, borrá los bloques de trabajo y poné una línea: {{EXPERIENCE_SECTION}}")
        return False

    # Leer estilo COMPLETO del bullet del template (fuente + pPr con sangría/numPr)
    ref_rpr, ref_ppr = _get_ref_bullet_info(doc)

    parent = placeholder_para._element.getparent()
    insert_idx = list(parent).index(placeholder_para._element)
    parent.remove(placeholder_para._element)

    new_elements = []
    for block_idx, block in enumerate(experience_blocks):
        title = block.get("title", "").strip()
        company_line = block.get("company_line", "").strip()
        bullets = block.get("bullets", [])

        # Título: 11pt, negrita, negro — sin herencia de estilo de lista
        before = 120 if block_idx > 0 else 0
        if title:
            new_elements.append(_make_paragraph(
                title, bold=True, ref_rpr=ref_rpr,
                before_twips=before, after_twips=0,
                font_size_half_pts=22, color_hex="000000"
            ))

        # Empresa/fecha: 10pt, itálica, gris — sin herencia de estilo de lista
        if company_line:
            new_elements.append(_make_paragraph(
                company_line, italic=True, ref_rpr=ref_rpr,
                before_twips=0, after_twips=0,
                font_size_half_pts=20, color_hex="595959"
            ))

        # Viñetas: heredan pPr COMPLETO del template (sangría idéntica a LANGUAGES)
        for bullet in bullets:
            if bullet.strip():
                new_elements.append(_make_paragraph(
                    bullet.strip(), ref_rpr=ref_rpr, ref_ppr=ref_ppr,
                    is_bullet=True,
                    before_twips=0, after_twips=0,
                    font_size_half_pts=20, color_hex="000000"
                ))

    for i, elem in enumerate(new_elements):
        parent.insert(insert_idx + i, elem)

    print(f" -> {len(experience_blocks)} bloque(s) de experiencia inyectados.")
    return True


def inject_skills_section(doc, skills_lines):
    """Busca {{SKILLS_SECTION}} y reemplaza con líneas de habilidades."""
    placeholder_para = _find_placeholder_para(doc, "{{SKILLS_SECTION}}")
    if placeholder_para is None:
        return False

    ref_rpr, ref_ppr = _get_ref_bullet_info(doc)
    parent = placeholder_para._element.getparent()
    insert_idx = list(parent).index(placeholder_para._element)
    parent.remove(placeholder_para._element)

    for i, line in enumerate(skills_lines):
        line = line.strip()
        if not line:
            continue
        elem = _make_paragraph(
            line, ref_rpr=ref_rpr, ref_ppr=ref_ppr,
            is_bullet=True, before_twips=0, after_twips=0,
            font_size_half_pts=20, color_hex="000000"
        )
        parent.insert(insert_idx + i, elem)

    print(f" -> {len(skills_lines)} línea(s) de habilidades inyectadas.")
    return True


def inject_education_section(doc, education_lines):
    """Busca {{EDUCATION_SECTION}} y reemplaza con líneas de educación."""
    placeholder_para = _find_placeholder_para(doc, "{{EDUCATION_SECTION}}")
    if placeholder_para is None:
        return False

    ref_rpr, ref_ppr = _get_ref_bullet_info(doc)
    parent = placeholder_para._element.getparent()
    insert_idx = list(parent).index(placeholder_para._element)
    parent.remove(placeholder_para._element)

    for i, line in enumerate(education_lines):
        line = line.strip()
        if not line:
            continue
        elem = _make_paragraph(
            line, ref_rpr=ref_rpr, ref_ppr=ref_ppr,
            is_bullet=True, before_twips=0, after_twips=0,
            font_size_half_pts=20, color_hex="000000"
        )
        parent.insert(insert_idx + i, elem)

    print(f" -> {len(education_lines)} línea(s) de educación inyectadas.")
    return True


# ─────────────────────────────────────────────────────────────────────────────
# REEMPLAZO DE VIÑETA INDIVIDUAL (para skills y education)
# ─────────────────────────────────────────────────────────────────────────────

def replace_bullet_in_docx(doc, original_clean, new_text):
    """
    Busca el párrafo cuyo texto limpio coincide y lo reemplaza preservando formato.
    """
    for para in _all_paragraphs(doc):
        p_clean = clean_text(para.text)
        if not p_clean or len(p_clean) < 5:
            continue
        if p_clean == original_clean or original_clean in p_clean or p_clean in original_clean:
            prefix = ""
            for char in para.text:
                if char in ["●", "•", "-", "*", " ", "\t"]:
                    prefix += char
                else:
                    break
            full_new = prefix + new_text

            if para.runs:
                first = para.runs[0]
                saved = {
                    'font_name': first.font.name,
                    'font_size': first.font.size,
                    'bold': first.bold,
                    'italic': first.italic,
                }
                try:
                    saved['color'] = first.font.color.rgb
                except Exception:
                    saved['color'] = None

                first.text = full_new
                for run in para.runs[1:]:
                    run.text = ""

                first.font.name = saved['font_name']
                first.font.size = saved['font_size']
                first.bold = saved['bold']
                first.italic = saved['italic']
                if saved['color']:
                    try:
                        first.font.color.rgb = saved['color']
                    except Exception:
                        pass
            else:
                para.text = full_new
            return True
    return False


# ─────────────────────────────────────────────────────────────────────────────
# FUNCIÓN PRINCIPAL: CONSTRUIR EL RESUME DESDE LA PLANTILLA
# ─────────────────────────────────────────────────────────────────────────────

def build_resume_from_template(tailored_data, output_path, template_path="base_resume.docx"):
    if not os.path.exists(template_path):
        print(f"\n[!] Advertencia: No se encontró la plantilla '{template_path}'.")
        return False

    try:
        doc = docx.Document(template_path)

        # 1. Perfil profesional (placeholder {{PROFESSIONAL_PROFILE}})
        profile_text = tailored_data.get("profile_tailored", "")
        if not profile_text or profile_text.startswith("Error"):
            profile_text = DEFAULT_PROFILE
        replace_placeholder(doc, "{{PROFESSIONAL_PROFILE}}", profile_text)

        # 2. Inyectar experiencia dinámica
        experience_blocks = tailored_data.get("experience_blocks", [])
        if not experience_blocks:
            experience_blocks = DEFAULT_EXPERIENCE_BLOCKS
        inject_experience_section(doc, experience_blocks)

        # 3. Habilidades clave
        skills_lines = tailored_data.get("skills_lines", [])
        if not skills_lines:
            skills_lines = DEFAULT_SKILLS_LINES
        inject_skills_section(doc, skills_lines)

        # 4. Educación
        education_lines = tailored_data.get("education_lines", [])
        if not education_lines:
            education_lines = DEFAULT_EDUCATION_LINES
        inject_education_section(doc, education_lines)

        # Limpiar cualquier placeholder residual si quedara alguno
        replace_placeholder(doc, "{{EXPERIENCE_SECTION}}", "")
        replace_placeholder(doc, "{{SKILLS_SECTION}}", "")
        replace_placeholder(doc, "{{EDUCATION_SECTION}}", "")

        doc.save(output_path)
        print(f"CV adaptado guardado con éxito en: {output_path}")
        return True
    except Exception as e:
        print(f"Error procesando la plantilla de CV: {e}")
        import traceback
        traceback.print_exc()
        return False


# ─────────────────────────────────────────────────────────────────────────────
# COVER LETTER
# ─────────────────────────────────────────────────────────────────────────────

def build_cover_letter_from_template(body_text, output_path, template_path="base_cover_letter.docx"):
    if os.path.exists(template_path):
        try:
            doc = docx.Document(template_path)
            replace_placeholder(doc, "{{COVER_LETTER_BODY}}", body_text)
            doc.save(output_path)
            print(f"Cover Letter guardada (usando plantilla): {output_path}")
            return True
        except Exception as e:
            print(f"Error procesando la plantilla de Cover Letter: {e}. Usando fallback.")

    try:
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = docx.shared.Pt(11)
        for line in body_text.split('\n'):
            doc.add_paragraph(line)
        doc.save(output_path)
        print(f"Cover Letter guardada (formato básico): {output_path}")
        return True
    except Exception as e:
        print(f"Error al guardar Cover Letter básica: {e}")
        return False


# ─────────────────────────────────────────────────────────────────────────────
# CONVERSIÓN A PDF
# ─────────────────────────────────────────────────────────────────────────────

def convert_to_pdf(docx_path, output_dir):
    print(f"Convirtiendo {os.path.basename(docx_path)} a PDF...")
    try:
        cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path]
        res = subprocess.run(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        if res.returncode == 0:
            pdf_path = docx_path.replace(".docx", ".pdf")
            print(f" -> PDF generado: {os.path.basename(pdf_path)}")
            return pdf_path
        else:
            print(f" -> LibreOffice no pudo convertir a PDF (código {res.returncode}). Se conserva el archivo .docx")
            return None
    except FileNotFoundError:
        print("\n[!] LibreOffice no instalado. Instalalo en WSL con:")
        print("  sudo apt update && sudo apt install -y libreoffice-nogui")
        return None
    except Exception as e:
        print(f" -> Error en conversión PDF: {e}")
        return None
